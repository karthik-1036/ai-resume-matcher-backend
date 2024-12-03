import PyPDF2
import docx
import io
import re
import logging

logger = logging.getLogger(__name__)

def clean_text(text):
    """Clean and normalize extracted text"""
    # Remove extra whitespace and normalize line endings
    text = re.sub(r'\s+', ' ', text)
    text = text.replace('\n\n', '\n').replace('\r\n', '\n')
    
    # Remove special characters but keep important punctuation
    text = re.sub(r'[^\w\s\.,;:\-\(\)@]', '', text)
    
    # Normalize common section headers
    text = re.sub(r'EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT|WORK HISTORY', 'EXPERIENCE:', text, flags=re.IGNORECASE)
    text = re.sub(r'EDUCATION|ACADEMIC|ACADEMICS', 'EDUCATION:', text, flags=re.IGNORECASE)
    text = re.sub(r'SKILLS|TECHNICAL SKILLS|EXPERTISE', 'SKILLS:', text, flags=re.IGNORECASE)
    
    return text.strip()

def extract_sections(text):
    """Extract and organize resume sections"""
    sections = {
        'experience': '',
        'education': '',
        'skills': '',
        'other': ''
    }
    
    # Split into sections based on common headers
    parts = re.split(r'(EXPERIENCE:|EDUCATION:|SKILLS:)', text, flags=re.IGNORECASE)
    
    current_section = 'other'
    for part in parts:
        part = part.strip()
        if not part:
            continue
            
        if 'EXPERIENCE:' in part.upper():
            current_section = 'experience'
        elif 'EDUCATION:' in part.upper():
            current_section = 'education'
        elif 'SKILLS:' in part.upper():
            current_section = 'skills'
        else:
            sections[current_section] += part + '\n'
    
    return sections

def parse_pdf(file_stream):
    """Extract text from PDF file"""
    try:
        logger.debug("Parsing PDF file")
        pdf_reader = PyPDF2.PdfReader(file_stream)
        text = ""
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        if not text.strip():
            logger.error("No text content found in PDF")
            raise ValueError("Could not extract text from PDF file")
            
        logger.debug("Successfully extracted text from PDF")
        return clean_text(text.strip())
        
    except Exception as e:
        logger.error(f"Error parsing PDF: {str(e)}")
        raise

def parse_docx(file_stream):
    """Extract text from DOCX file"""
    try:
        logger.debug("Parsing DOCX file")
        doc = docx.Document(file_stream)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
                
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
                        
        if not text.strip():
            logger.error("No text content found in DOCX")
            raise ValueError("Could not extract text from DOCX file")
            
        logger.debug("Successfully extracted text from DOCX")
        return clean_text(text.strip())
        
    except Exception as e:
        logger.error(f"Error parsing DOCX: {str(e)}")
        raise

def parse_resume(file):
    """
    Parse resume file (PDF or DOCX) and extract text content
    """
    try:
        filename = file.filename.lower()
        logger.debug(f"Parsing resume file: {filename}")
        
        # Create a file stream
        file_stream = io.BytesIO(file.read())
        
        if filename.endswith('.pdf'):
            text = parse_pdf(file_stream)
        elif filename.endswith('.docx'):
            text = parse_docx(file_stream)
        else:
            logger.error(f"Unsupported file format: {filename}")
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
            
        # Extract and organize sections
        sections = extract_sections(text)
        
        # Combine sections in a structured way
        structured_text = f"""
EXPERIENCE:
{sections['experience']}

EDUCATION:
{sections['education']}

SKILLS:
{sections['skills']}

ADDITIONAL INFORMATION:
{sections['other']}
"""
        return structured_text.strip()
        
    except Exception as e:
        logger.error(f"Error parsing resume: {str(e)}")
        raise
